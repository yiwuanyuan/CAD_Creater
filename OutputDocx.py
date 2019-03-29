#!/usr/bin/env python
# -*- coding: utf-8 -*-
# File  : OutputDocx.py
# Author: Wangyuan
# Date  : 2019-1-3

from docx import Document
from docx.shared import Mm
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import math
import time
from os import chdir


def size_info(parameter):
    if 'od' in parameter and not ('id' in parameter):
        size = '直径: ' + str(parameter['od'])
        return size

    elif 'od' in parameter and 'id' in parameter and not ('angle' in parameter):
        size = '外直径: ' + str(parameter['od']) + ' 内直径: ' + str(parameter['id'])
        return size

    elif 'od' in parameter and 'id' in parameter and 'angle' in parameter:
        size = '外直径: ' + str(parameter['od']) + ' 内直径: ' + str(parameter['id']) + ' 弧度: ' + str(parameter['angle'])
        return size
    elif 'w' in parameter and 'l' in parameter:
        if int(parameter['l']) > parameter['w']:
            size = '长度: ' + str(math.ceil(parameter['l'])) + ' 宽度: ' + str(math.ceil(parameter['w']))
        else:
            size = '长度: ' + str(math.ceil(parameter['w'])) + ' 宽度: ' + str(math.ceil(parameter['l']))
        return size

    else:
        size = '错误内容'
        return size


# 改写成一个类
def OutputDocx(info, address, isNeedML):
    material_gather = []
    addr = '/'.join(address.strip().strip('\'').split('/')[:-1])
    chdir(addr)
    # doc = Document(docx=path.join(getcwd(), 'default.docx'))
    doc = Document()
    if address.find('排料清单') != -1:
        table_title = str(address.split('\\')[-1].split('/')[-1].split('_')[0].split('-')[0].split(' ')[0])
    else:
        table_title = str(address.split('\\')[-1].split('/')[-1].split(' ')[0])

    # 设置全局字体
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # 设置word文件的页面属性
    sections = doc.sections
    sections[0].page_width = Mm(297)
    sections[0].page_height = Mm(210)
    # 添加基本内容
    tit = doc.add_paragraph('结构件下料指导书')
    # 新设置一种名为"title_style"的新style,设置字体大小，字体样式
    title_style = doc.styles.add_style('UserStyle1', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(22)
    title_style.font.bold = True
    title_style.font.name = u'宋体'
    title_style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    tit.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 为标题应用样式
    tit.style = title_style

    # 添加内容，绘制表头
    title_list = [r'序号', r'产品代号', r'零件名称', r'材料', u'厚度', u'下料净尺寸(mm)', u'数量', r'备注']
    trow = len(info) + 2
    tcol = len(title_list)
    table = doc.add_table(rows=trow, cols=tcol, style='Table Grid')
    table.cell(0, 0).merge(table.cell(0, tcol - 1))  # 合并第一行

    # 新设置一种名为"Bold_style"的新style,设置字体大小，字体样式
    Bold_style = doc.styles.add_style('UserStyle2', WD_STYLE_TYPE.CHARACTER)
    Bold_style.font.size = Pt(14)
    Bold_style.font.bold = True
    Bold_style.font.name = u'宋体'
    Bold_style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # 为表注应用样式
    run = table.cell(0, 0).paragraphs[0].add_run(
        "  合同号：%s" % table_title + "                     文件编号：JXL-%s-001" % table_title)
    table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run.style = Bold_style

    # 设置表格内的字体
    table.style.font.size = Pt(12)
    table.style.font.name = u'微软雅黑'
    table.style._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    # 设置表格对其方式
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # WD_TABLE_ALIGNMENT.LEFT|WD_TABLE_ALIGNMENT.RIGHT 其他设置方式

    # 新设置一种名为"tbhead"的新style,设置字体大小，字体样式
    tbhead = doc.styles.add_style('UserStyle3', WD_STYLE_TYPE.CHARACTER)
    tbhead.font.size = Pt(12)
    tbhead.font.bold = True
    tbhead.font.name = u'宋体'
    tbhead._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # 设置表格的表头的样式为table_head
    for i, value in enumerate(title_list):
        run_1 = table.cell(1, i).paragraphs[0].add_run(value)
        table.cell(1, i).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run_1.style = tbhead
    # 设置表格内容及对其方式
    y = 0
    for list in info:
        item_material = {'material':str(list['material']),'thickness':str(list['thickness'])}
        # print(item_material)
        if item_material not in material_gather:
            material_gather.append(item_material)
        x = 1
        # 设置每一行第一个单元格的数值
        table.cell(y + 2, 0).text = str(y + 1)
        table_result = '数据已处理: ' + str(int((y + 1) / len(info) * 100)) + '%'
        print(table_result)

        table.cell(y + 2, 0).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(y + 2, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 设置每一行的单元格的内容
        table.cell(y + 2, 1).text = str(list['pro_name'])
        table.cell(y + 2, 2).text = str(list['part_name'])
        table.cell(y + 2, 3).text = str(list['material'])
        table.cell(y + 2, 4).text = str(list['thickness'])
        table.cell(y + 2, 5).text = size_info(list['parameter'])
        table.cell(y + 2, 6).text = str(list['sum'])
        if list['remark'] == list['remark']:
            table.cell(y + 2, 7).text = str(list['remark'])
        for x in range(tcol):
            table.cell(y + 2, x).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(y + 2, x).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        y += 1
    print('数据已处理完毕，请耐心等待')
    # 设置列宽
    table.autofit = False
    for i in range(trow):
        table.cell(i, 0).width = Mm(15)
        table.cell(i, 1).width = Mm(45)
        table.cell(i, 2).width = Mm(30)
        table.cell(i, 3).width = Mm(30)
        table.cell(i, 4).width = Mm(20)
        table.cell(i, 5).width = Mm(65)
        table.cell(i, 6).width = Mm(14)
        table.cell(i, 7).width = Mm(26)
        # 设置行高
        table.rows[i].height = Mm(11)

    # 设置表格第一行的对齐方式为左对齐
    table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    #材料清单相关：
    if isNeedML:
        material_mark_title = [r'序号', r'材料牌号', r'厚度', r'版幅及数量', r'备注']
        material_mark_rows=len(material_gather)+2
        material_mark = doc.add_table(material_mark_rows, cols=len(material_mark_title), style='Table Grid') #创建表格
        material_mark.autofit = False
        for k in range(material_mark_rows):
            material_mark.cell(k, 0).width = Mm(15)
            material_mark.cell(k, 1).width = Mm(30)
            material_mark.cell(k, 2).width = Mm(30)
            material_mark.cell(k, 3).width = Mm(125)
            material_mark.cell(k, 4).width = Mm(40)

        material_mark.alignment = WD_TABLE_ALIGNMENT.CENTER #表格居中
        material_mark.cell(0, 0).merge(material_mark.cell(0, (len(material_mark_title)-1))) #合并第一行
        material_mark.cell(0, 0).paragraphs[0].add_run('以上零件所用板材清单：')
        material_mark_title = [r'序号', r'材料牌号', r'厚度', r'版幅及数量', r'备注']
        #表头添加
        for i in range(len(material_mark_title)):
            material_mark.cell(1, i).paragraphs[0].add_run(material_mark_title[i]).style = tbhead
            material_mark.cell(1, i).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            material_mark.cell(1, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        #内容样式添加
        material_mark_no = 0
        for j in range(material_mark_rows-2):
            material_mark.cell(j+2,0).paragraphs[0].add_run(str(material_mark_no))
            material_mark.cell(j+2,1).paragraphs[0].add_run(material_gather[j]['material'])
            material_mark.cell(j+2,2).paragraphs[0].add_run(material_gather[j]['thickness'])
            material_mark_no += 1

            # 材料清单栏样式调节
            material_mark.rows[j+2].height = Mm(11)



        material_mark.rows[0].height = Mm(12)
        material_mark.rows[1].height = Mm(12)
        print(material_gather)

        # for i in material_gather:
        #     material_mark2 = par1.add_run('\n        ' + i + ' 版幅及数量:')
        #     material_mark1.style = Bold_style
        #     material_mark2.style = Bold_style


    # 新备注内容
    table_mark = doc.add_table(rows=1, cols=1, style='Table Grid')
    table_mark.alignment = WD_TABLE_ALIGNMENT.CENTER

    # table_mark.cell(0, 0).merge(table.cell(0, tcol - 1))  # 合并第一行
    mark_content = table_mark.cell(0, 0).paragraphs[0].add_run('注意事项:' +
                                                               u'\n        1.完成后请做好标识，标识内容包括：合同号、产品代号、零件号、尺寸。' +
                                                               u'\n        2.大尺寸余料请进行尺寸测量，记录过后请注意保管。')
    table_mark.rows[0].height = Mm(25)
    table_mark.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(6)
    mark_content.style = Bold_style


    # 签名栏重写
    table_sign = doc.add_table(rows=1, cols=6, style='Table Grid')
    table_sign.alignment = WD_TABLE_ALIGNMENT.CENTER

    sign_style = doc.styles.add_style('UserStyle4', WD_STYLE_TYPE.CHARACTER)
    sign_style.font.size = Pt(14)
    sign_style.font.bold = True
    sign_style.font.name = u'宋体'
    sign_style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # table_sign.cell(0, 0).merge(table.cell(0, tcol - 1))  # 合并第一行
    table_sign.cell(0, 0).paragraphs[0].add_run('编 制:').style = sign_style
    table_sign.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_sign.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table_sign.cell(0, 2).paragraphs[0].add_run('校 对:').style = sign_style
    table_sign.cell(0, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_sign.cell(0, 2).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table_sign.cell(0, 4).paragraphs[0].add_run('调度员:').style = sign_style
    table_sign.cell(0, 4).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_sign.cell(0, 4).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER



    # 备注栏样式修改
    table_sign.cell(0, 0).width = Mm(30)
    table_sign.cell(0, 1).width = Mm(50)
    table_sign.cell(0, 2).width = Mm(30)
    table_sign.cell(0, 3).width = Mm(50)
    table_sign.cell(0, 4).width = Mm(30)
    table_sign.cell(0, 5).width = Mm(50)

    table_sign.rows[0].height = Mm(12)

    # 旧备注内容
    par1 = doc.add_paragraph('')
    par2 = doc.add_paragraph('')
    par3 = doc.add_paragraph('')
    par1.paragraph_format.space_after = 0
    par2.paragraph_format.space_after = 0
    par2.paragraph_format.space_before = Pt(6)
    par3.paragraph_format.space_before = Pt(2)
    mark1 = par1.add_run(u'\n  备注:')
    mark2 = par1.add_run(
        u'\n        完成后请做好标识，标识内容包括：合同号、产品代号、零件号、尺寸，余料请注意带回。')
    # if isNeedML:
    #     material_mark1 = par1.add_run(u'\n        以下为所用板材清单：')
    #     for i in material_gather:
    #         material_mark2 = par1.add_run('\n        ' + i + ' 版幅及数量:')
    #         material_mark1.style = Bold_style
    #         material_mark2.style = Bold_style
    # # mark3 = par2.add_run(u'           编制:王元 %s。'%time.strftime('%Y-%m-%d',time.localtime(time.time()))
    # #                     +'                  审核:刘光 %s。'%time.strftime('%Y-%m-%d',time.localtime(time.time())))
    mark3 = par2.add_run('  编 制:         时 间:       '
                         + '      审 核:         时 间:       ')
    mark4 = par2.add_run('      生产确认:         时 间:       ')
    mark1.style = Bold_style
    mark2.style = Bold_style
    mark3.style = Bold_style
    mark4.style = Bold_style
    # 设定页边距为上18mm.下15mm
    doc.sections[0].top_margin = Mm(18)
    doc.sections[0].bottom_margin = Mm(15)
    # 文件保存
    print(addr + '\\' + table_title + u'_结构件下料清单.docx')
    doc.save(addr + '\\' + table_title + u'_结构件下料清单.docx')
    table_finish = '文件已生成'
    print(table_finish)


if __name__ == '__main__':
    from GetExcelInfo import *

    addr = '3032017XXXX_排料清单.xlsx'
    g = GetExcelInfo(addr)
    info = g.output
    OutputDocx(info, 'C:/Users/wangyuan/PycharmProjects/CAD_Creater/' + addr, True)
